"""
Generates the Internship Report Word Document
following university guidelines.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color)
    shading_elm.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading_elm)


def create_document():
    doc = Document()

    # Set margins: 2.00 cm all sides
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Configure styles
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.5

    # ================================================================
    # INNER TITLE PAGE
    # ================================================================
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INTERNSHIP REPORT")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("on")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    project_title = doc.add_paragraph()
    project_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = project_title.add_run("Contract Testing Automation Using AI-Driven Agent")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    doc.add_paragraph()
    doc.add_paragraph()

    at_para = doc.add_paragraph()
    at_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = at_para.add_run("Carried out at")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    company_para = doc.add_paragraph()
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company_para.add_run("Bottomline Technologies")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()
    doc.add_paragraph()

    by_para = doc.add_paragraph()
    by_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = by_para.add_run("Submitted by")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run("Druva S Kumar")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    # University details placeholder
    uni_para = doc.add_paragraph()
    uni_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = uni_para.add_run("[University Name]\n[Department]\n[Registration Number]")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_page_break()

    # ================================================================
    # COMPANY CERTIFICATE PAGE (Placeholder)
    # ================================================================
    for _ in range(3):
        doc.add_paragraph()

    cert_title = doc.add_paragraph()
    cert_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cert_title.add_run("COMPANY CERTIFICATE")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    doc.add_paragraph()
    doc.add_paragraph()

    cert_note = doc.add_paragraph()
    cert_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cert_note.add_run("[Attach Company Completion Certificate or Acceptance/Joining Letter here]")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.italic = True

    doc.add_page_break()

    # ================================================================
    # DECLARATION
    # ================================================================
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run("DECLARATION")
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    decl_text = (
        "I, Druva S Kumar, hereby declare that the Internship Report entitled "
        '"Contract Testing Automation Using AI-Driven Agent" submitted to '
        "[University Name] is a record of an original work done by me during my "
        "internship at Bottomline Technologies, and this work has not been submitted "
        "elsewhere for any other degree or diploma."
    )
    p = doc.add_paragraph(decl_text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Place: [City]\nDate: [Date]\n\n\n\nDruva S Kumar\n[Registration Number]")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_page_break()

    # ================================================================
    # ACKNOWLEDGEMENT
    # ================================================================
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run("ACKNOWLEDGEMENT")
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    ack_text = (
        "I would like to express my sincere gratitude to Bottomline Technologies for providing "
        "me the opportunity to undertake this internship. I am deeply thankful to my industry "
        "mentor and the engineering team for their guidance, support, and valuable insights "
        "throughout the duration of this project."
    )
    p = doc.add_paragraph(ack_text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ack_text2 = (
        "I would also like to thank my university faculty and internship coordinator for "
        "their continuous encouragement and for providing the academic framework that made "
        "this internship possible."
    )
    p = doc.add_paragraph(ack_text2)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ack_text3 = (
        "Finally, I extend my appreciation to my peers and family members for their "
        "unwavering support during this period."
    )
    p = doc.add_paragraph(ack_text3)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()

    # ================================================================
    # ABSTRACT
    # ================================================================
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run("ABSTRACT")
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    abstract_text = (
        "In modern microservices architectures, ensuring API compatibility between services "
        "is critical to prevent production failures. Contract testing provides a mechanism to "
        "verify that API providers and consumers adhere to agreed-upon contracts. However, "
        "maintaining these contracts manually as APIs evolve is error-prone and time-consuming."
    )
    p = doc.add_paragraph(abstract_text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    abstract_text2 = (
        "This internship project addresses this challenge by developing an AI-driven automation "
        "agent that manages the complete lifecycle of contract testing. The system uses Spring "
        "Cloud Contract for contract verification between a Provider API (User Service) and a "
        "Consumer API (Order Service), while a Python-based AI agent automates contract generation, "
        "drift detection, auto-remediation, CI/CD pipeline integration, and team notifications."
    )
    p = doc.add_paragraph(abstract_text2)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    abstract_text3 = (
        "The project delivers a fully automated pipeline where contract drift is detected "
        "in real-time, fixes are generated and submitted as merge requests automatically, "
        "and developers are notified via email when issues are found. A web-based dashboard "
        "provides visibility into contract health across the system. This eliminates manual "
        "intervention in contract maintenance and reduces the risk of API-breaking changes "
        "reaching production environments."
    )
    p = doc.add_paragraph(abstract_text3)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()

    # ================================================================
    # TABLE OF CONTENTS (Placeholder - Word can auto-generate)
    # ================================================================
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run("TABLE OF CONTENTS")
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    toc_items = [
        ("1.", "Introduction", ""),
        ("2.", "Company Introduction", ""),
        ("3.", "Internship Project Details", ""),
        ("3.1", "Project Title", ""),
        ("3.2", "Project Context", ""),
        ("3.3", "Roles and Responsibilities", ""),
        ("3.4", "Project Abstract and Scope", ""),
        ("3.5", "Project Design and Technologies Used", ""),
        ("3.6", "Implementation Details", ""),
        ("3.7", "Project Results and Learning Outcomes", ""),
        ("4.", "Conclusion", ""),
        ("5.", "References", ""),
    ]

    for num, title_text, _ in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}\t{title_text}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.add_page_break()

    # ================================================================
    # LIST OF TABLES AND FIGURES
    # ================================================================
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run("LIST OF TABLES AND FIGURES")
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run("List of Tables")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    tables_list = [
        "Table 1: Technologies Used in the Project",
        "Table 2: CI/CD Pipeline Stages and Jobs",
        "Table 3: AI Agent CLI Commands",
        "Table 4: Notification Scenarios",
    ]
    for t in tables_list:
        p = doc.add_paragraph(t)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run("List of Figures")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    figures_list = [
        "Figure 1: High-Level System Architecture",
        "Figure 2: Contract Testing Flow",
        "Figure 3: CI/CD Pipeline Architecture",
        "Figure 4: AI Agent Module Architecture",
        "Figure 5: Contract Health Dashboard",
        "Figure 6: Drift Detection and Auto-Fix Workflow",
    ]
    for f in figures_list:
        p = doc.add_paragraph(f)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_page_break()

    # ================================================================
    # CHAPTER 1: INTRODUCTION
    # ================================================================
    add_heading_formatted(doc, "1. Introduction")

    intro_text = (
        "Application Programming Interfaces (APIs) form the backbone of modern software systems, "
        "enabling communication between microservices, mobile applications, and third-party integrations. "
        "As organizations adopt microservices architectures, the number of inter-service API dependencies "
        "grows exponentially. A single breaking change in one service's API can cascade into failures "
        "across multiple dependent services, leading to production outages and degraded user experience."
    )
    p = doc.add_paragraph(intro_text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    intro_text2 = (
        "Contract testing is a software testing methodology that addresses this challenge by defining "
        "explicit agreements (contracts) between API providers and consumers. These contracts specify "
        "the expected request and response formats, allowing each service to be tested independently "
        "against the contract rather than requiring full end-to-end integration tests."
    )
    p = doc.add_paragraph(intro_text2)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    intro_text3 = (
        "However, maintaining contracts manually presents significant challenges. As APIs evolve, "
        "contracts can become outdated (drift), new endpoints may lack contracts (uncovered), and "
        "deleted endpoints may leave behind orphaned contracts. Detecting and fixing these issues "
        "manually is time-consuming and error-prone."
    )
    p = doc.add_paragraph(intro_text3)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    intro_text4 = (
        "This internship project solves these challenges by building an AI-driven automation agent "
        "that manages the complete contract testing lifecycle — from generating contracts based on "
        "OpenAPI specifications, to detecting drift, auto-fixing issues, integrating with CI/CD "
        "pipelines, and notifying developers when action is required. The goal is to make contract "
        "compliance fully automated and zero-maintenance."
    )
    p = doc.add_paragraph(intro_text4)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()

    # ================================================================
    # CHAPTER 2: COMPANY INTRODUCTION
    # ================================================================
    add_heading_formatted(doc, "2. Company Introduction")

    company_text = (
        "Bottomline Technologies is a global financial technology company that provides "
        "cloud-based solutions for business payments, financial document automation, and "
        "digital banking. Headquartered in the United States, the company serves financial "
        "institutions, corporations, and banks across the world, helping them modernize "
        "their payment and financial processes."
    )
    p = doc.add_paragraph(company_text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    company_text2 = (
        "The company's product portfolio includes solutions for payment processing, fraud "
        "detection, financial messaging, and regulatory compliance. Bottomline's engineering "
        "teams follow modern software development practices including microservices architecture, "
        "continuous integration/continuous deployment (CI/CD), and cloud-native development."
    )
    p = doc.add_paragraph(company_text2)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    company_text3 = (
        "As part of the Core Services engineering division, the internship focused on improving "
        "the quality and reliability of inter-service API communication through automated "
        "contract testing. The team follows agile development methodologies with GitLab as the "
        "primary platform for source control, CI/CD, and collaboration."
    )
    p = doc.add_paragraph(company_text3)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()

    # ================================================================
    # CHAPTER 3: INTERNSHIP PROJECT DETAILS
    # ================================================================
    add_heading_formatted(doc, "3. Internship Project Details")

    # 3.1 Project Title
    add_subheading_formatted(doc, "3.1 Project Title")

    p = doc.add_paragraph("Contract Testing Automation Using AI-Driven Agent")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 3.2 Project Context
    add_subheading_formatted(doc, "3.2 Project Context")

    context_text = (
        "This project is part of the broader quality assurance and DevOps initiative within "
        "the engineering division. The organization's microservices ecosystem comprises multiple "
        "services that communicate via REST APIs. Ensuring API compatibility across services "
        "during rapid development cycles was identified as a key challenge. This project was "
        "initiated as an independent initiative to build tooling that automates contract "
        "testing and integrates it into the existing CI/CD workflow."
    )
    p = doc.add_paragraph(context_text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 3.3 Roles and Responsibilities
    add_subheading_formatted(doc, "3.3 Roles and Responsibilities")

    roles_text = (
        "As the sole developer on this project, the responsibilities encompassed the full "
        "software development lifecycle:"
    )
    p = doc.add_paragraph(roles_text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    roles = [
        "Designing the system architecture for contract testing automation",
        "Developing the Provider API (User Service) and Consumer API (Order Service) as demonstration microservices",
        "Implementing Spring Cloud Contract for provider-side contract verification and consumer-side stub testing",
        "Building the Python-based AI agent with modules for contract generation, drift detection, report generation, and auto-remediation",
        "Designing and implementing the GitLab CI/CD pipeline with automated contract testing gates",
        "Developing a web-based dashboard for real-time contract health monitoring",
        "Implementing automated merge request creation for contract fixes",
        "Building a notification system for developer alerts via email",
        "Writing documentation and conducting end-to-end testing",
    ]
    for role in roles:
        p = doc.add_paragraph(role, style="List Bullet")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 3.4 Project Abstract and Scope
    add_subheading_formatted(doc, "3.4 Project Abstract and Scope")

    scope_text = (
        "The project scope encompasses building a complete contract testing automation system "
        "consisting of three main components:"
    )
    p = doc.add_paragraph(scope_text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = doc.add_paragraph(
        "Provider API (User Service): A Spring Boot REST API that manages user data "
        "with CRUD operations. It exposes an OpenAPI specification that serves as the "
        "source of truth for contract generation.", style="List Number"
    )
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = doc.add_paragraph(
        "Consumer API (Order Service): A Spring Boot REST API that depends on the "
        "Provider API. It uses the UserServiceClient to fetch user information for "
        "order processing, demonstrating real-world inter-service dependencies.", style="List Number"
    )
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = doc.add_paragraph(
        "AI Agent: A Python-based command-line tool that automates the entire contract "
        "testing lifecycle — generation, validation, drift detection, auto-fix, CI/CD "
        "integration, dashboard, and notifications.", style="List Number"
    )
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()

    # 3.5 Project Design and Technologies
    add_subheading_formatted(doc, "3.5 Project Design and Technologies Used")

    design_text = (
        "The system follows a modular architecture where each component operates independently "
        "but integrates through well-defined interfaces. The design principles include:"
    )
    p = doc.add_paragraph(design_text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    principles = [
        "Separation of Concerns: Each module handles a single responsibility",
        "OpenAPI as Source of Truth: All contract generation derives from the live API specification",
        "Fail-Fast in CI/CD: Contract failures block deployment to prevent breaking changes",
        "Zero-Maintenance: Contracts are auto-generated and auto-fixed, requiring no manual intervention",
        "Observable: Dashboard and notifications provide visibility into system health",
    ]
    for pr in principles:
        p = doc.add_paragraph(pr, style="List Bullet")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    # Technologies Table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Technology"
    hdr_cells[1].text = "Version"
    hdr_cells[2].text = "Purpose"
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    tech_data = [
        ("Java", "22+", "Provider and Consumer API development"),
        ("Spring Boot", "3.2.5", "REST API framework with embedded server"),
        ("Spring Cloud Contract", "4.1.3", "Contract verification and stub generation"),
        ("Python", "3.x", "AI Agent development"),
        ("Flask", "3.x", "Contract Health Dashboard web UI"),
        ("Maven", "3.9", "Java project build and dependency management"),
        ("GitLab CI/CD", "-", "Continuous integration and deployment pipeline"),
        ("OpenAPI 3.0", "-", "API specification standard"),
        ("YAML", "-", "Contract file format"),
        ("SMTP", "-", "Email notification delivery"),
    ]
    for tech, ver, purpose in tech_data:
        row_cells = table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = ver
        row_cells[2].text = purpose

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Table 1: Technologies Used in the Project")
    run.italic = True
    run.font.size = Pt(10)

    doc.add_page_break()

    # 3.6 Implementation Details
    add_subheading_formatted(doc, "3.6 Implementation Details (Modules)")

    impl_text = (
        "The project was implemented in a phased approach, with each phase building upon "
        "the previous one. The system comprises the following key modules:"
    )
    p = doc.add_paragraph(impl_text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Module 1: Provider API
    add_sub_subheading(doc, "3.6.1 Provider API (User Service)")
    p = doc.add_paragraph(
        "The Provider API is a RESTful web service built with Spring Boot that manages "
        "user entities. It exposes five endpoints for CRUD operations (Create, Read, Update, "
        "Delete) on user resources. The API automatically generates an OpenAPI 3.0 specification "
        "via SpringDoc, which serves as the single source of truth for contract generation. "
        "Spring Cloud Contract Maven plugin reads YAML contract files from the test resources "
        "directory, auto-generates JUnit verification tests, and produces a stubs JAR for "
        "consumer-side testing."
    )
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Module 2: Consumer API
    add_sub_subheading(doc, "3.6.2 Consumer API (Order Service)")
    p = doc.add_paragraph(
        "The Consumer API demonstrates a real-world service dependency. It provides order "
        "management functionality and depends on the Provider API to fetch user information. "
        "The UserServiceClient makes HTTP calls to the Provider. During testing, Spring Cloud "
        "Contract Stub Runner replaces the real Provider with a WireMock server loaded with "
        "stubs generated from the contracts. This enables isolated testing of the Consumer "
        "without requiring the Provider to be running."
    )
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Module 3: AI Agent
    add_sub_subheading(doc, "3.6.3 AI Agent — Core Modules")
    p = doc.add_paragraph(
        "The AI Agent is the central automation tool, implemented in Python with a "
        "command-line interface. It consists of the following modules:"
    )
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    modules = [
        ("OpenAPI Spec Reader", "Fetches and parses the live OpenAPI specification from the running Provider API, extracting endpoint details including paths, methods, request/response schemas, and field types."),
        ("Contract Generator", "Generates Spring Cloud Contract YAML files from the parsed OpenAPI specification. It creates contracts with regex-based matchers for flexible validation rather than hardcoded values."),
        ("Drift Detector", "Compares existing contract files against the current OpenAPI specification to identify three types of issues: drifted contracts (schema mismatch), uncovered endpoints (no contract), and orphaned contracts (no matching endpoint)."),
        ("Report Generator", "Produces detailed reports showing contract health status, coverage percentage, and specific drift details with remediation suggestions."),
        ("CI Config Generator", "Generates a complete GitLab CI/CD pipeline configuration with multi-stage builds, contract tests, drift detection, auto-fix, and deployment gates."),
        ("MR Creator", "Automatically creates GitLab Merge Requests containing auto-fixed contracts, enabling a human-in-the-loop review process before changes are merged."),
        ("Notifier", "Sends automated notifications to developers via email when contract drift is detected or tests fail, using the corporate SMTP relay for delivery."),
    ]
    for name, desc in modules:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f"{name}: ")
        run.bold = True
        run.font.name = "Times New Roman"
        p.add_run(desc).font.name = "Times New Roman"

    doc.add_page_break()

    # Module 4: CI/CD Pipeline
    add_sub_subheading(doc, "3.6.4 CI/CD Pipeline")
    p = doc.add_paragraph(
        "The GitLab CI/CD pipeline enforces contract compliance on every code push. "
        "It is organized into five stages:"
    )
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Pipeline stages table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Stage"
    hdr_cells[1].text = "Jobs"
    hdr_cells[2].text = "Purpose"
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    pipeline_data = [
        ("Build", "provider-build, consumer-build", "Compile source code and detect compilation errors"),
        ("Test", "provider-contract-test, consumer-contract-test, ai-agent-drift-check", "Run contract verification, stub testing, and drift detection"),
        ("Report", "contract-report, notify-team", "Generate test reports and send notifications"),
        ("Fix", "auto-fix-contracts", "Auto-remediate drifted contracts (manual trigger)"),
        ("Deploy", "deploy", "Gated deployment — only if all tests pass"),
    ]
    for stage, jobs, purpose in pipeline_data:
        row_cells = table.add_row().cells
        row_cells[0].text = stage
        row_cells[1].text = jobs
        row_cells[2].text = purpose

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Table 2: CI/CD Pipeline Stages and Jobs")
    run.italic = True
    run.font.size = Pt(10)

    doc.add_paragraph()

    p = doc.add_paragraph(
        "The pipeline uses a deployment gate pattern — contract test failures block "
        "deployment, ensuring that API-breaking changes cannot reach production. The "
        "auto-fix job provides a one-click remediation option that generates correct "
        "contracts and submits them as a merge request for review."
    )
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Module 5: Dashboard
    add_sub_subheading(doc, "3.6.5 Contract Health Dashboard")
    p = doc.add_paragraph(
        "A Flask-based web dashboard provides real-time visibility into contract health. "
        "It displays the overall health status (Healthy, Warning, Critical), coverage "
        "percentage, per-endpoint breakdown, drift detection results, and historical "
        "health trends. The dashboard fetches live data from the running Provider API "
        "and presents it in an intuitive web interface accessible at a configurable port."
    )
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Module 6: Notifications
    add_sub_subheading(doc, "3.6.6 Notification System")
    p = doc.add_paragraph(
        "The notification module closes the feedback loop by alerting developers "
        "immediately when issues are detected. It supports email delivery via corporate "
        "SMTP relay and Slack webhooks. Notifications are triggered automatically by the "
        "CI/CD pipeline and include details about the drift, coverage metrics, and "
        "actionable remediation steps. When the SMTP server is not reachable (e.g., during "
        "local development), the system falls back to saving notification previews locally."
    )
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # CLI commands table
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Command"
    hdr_cells[1].text = "Function"
    hdr_cells[2].text = "Auto-Notify"
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    cmd_data = [
        ("generate", "Generate contract YAML files from OpenAPI spec", "No"),
        ("drift", "Detect drift between contracts and live spec", "Yes"),
        ("report", "Full contract health report", "Yes"),
        ("validate", "Validate contracts for CI/CD", "Yes"),
        ("fix", "Auto-fix drifted contracts + create merge request", "Yes"),
        ("notify", "Send standalone notification", "Always"),
        ("ci", "Generate GitLab CI/CD pipeline configuration", "No"),
        ("dashboard", "Start web-based health dashboard", "No"),
    ]
    for cmd, func, notify in cmd_data:
        row_cells = table.add_row().cells
        row_cells[0].text = cmd
        row_cells[1].text = func
        row_cells[2].text = notify

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Table 3: AI Agent CLI Commands")
    run.italic = True
    run.font.size = Pt(10)

    doc.add_page_break()

    # 3.7 Project Results and Learning Outcomes
    add_subheading_formatted(doc, "3.7 Project Results and Learning Outcomes")

    add_sub_subheading(doc, "3.7.1 Results")

    results_text = (
        "The project successfully achieved its objectives of automating the complete "
        "contract testing lifecycle. Key measurable outcomes include:"
    )
    p = doc.add_paragraph(results_text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    results = [
        "100% contract coverage achieved across all five API endpoints",
        "Automated drift detection with three severity levels (Healthy, Warning, Critical)",
        "End-to-end automated workflow: drift detection to fix to merge request creation in a single pipeline run",
        "Zero manual intervention required for contract maintenance during normal API evolution",
        "Real-time developer notifications enabling rapid response to breaking changes",
        "Web dashboard providing at-a-glance visibility into API contract health",
        "Complete CI/CD integration with deployment gates preventing broken APIs from reaching production",
    ]
    for r in results:
        p = doc.add_paragraph(r, style="List Bullet")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Notification scenarios table
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Scenario"
    hdr_cells[1].text = "Action"
    hdr_cells[2].text = "Notification"
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    notif_data = [
        ("Drift detected (Critical)", "Auto-fix + MR creation", "Email with drift details"),
        ("Drift detected (Warning)", "Alert developer", "Email with coverage report"),
        ("All contracts healthy", "No action needed", "Confirmation email"),
        ("Contract test failure", "Block deployment", "Email with failure details"),
        ("Auto-fix completed", "MR ready for review", "Email with MR link"),
    ]
    for scenario, action, notif in notif_data:
        row_cells = table.add_row().cells
        row_cells[0].text = scenario
        row_cells[1].text = action
        row_cells[2].text = notif

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Table 4: Notification Scenarios")
    run.italic = True
    run.font.size = Pt(10)

    doc.add_paragraph()

    add_sub_subheading(doc, "3.7.2 Learning Outcomes")

    learnings = [
        ("Spring Cloud Contract Framework", "Gained hands-on experience with consumer-driven contract testing, understanding the provider verification and consumer stub testing paradigms."),
        ("Microservices Architecture", "Learned how services communicate via REST APIs and the challenges of maintaining compatibility across service boundaries."),
        ("CI/CD Pipeline Design", "Designed and implemented a multi-stage GitLab CI/CD pipeline with dependency management, artifact passing, and deployment gates."),
        ("Python CLI Tool Development", "Built a modular command-line application with argparse, multiple subcommands, and extensible architecture."),
        ("API Specification Standards", "Worked extensively with OpenAPI 3.0 specifications and learned to parse and interpret schema definitions programmatically."),
        ("Git Workflow", "Practiced feature branching, merge requests, code review workflows, and automated branch management via APIs."),
        ("DevOps Practices", "Understood the importance of automated quality gates, fast feedback loops, and developer experience in modern software delivery."),
        ("SMTP and Notification Systems", "Learned about email delivery mechanisms, corporate SMTP relays, and building reliable notification pipelines."),
    ]
    for title_text, desc in learnings:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        run.font.name = "Times New Roman"
        p.add_run(desc).font.name = "Times New Roman"

    doc.add_page_break()

    # ================================================================
    # CHAPTER 4: CONCLUSION
    # ================================================================
    add_heading_formatted(doc, "4. Conclusion")

    conclusion_text = (
        "This internship project successfully demonstrates that contract testing can be fully "
        "automated using an AI-driven agent, eliminating the manual overhead traditionally "
        "associated with maintaining API contracts in microservices architectures. The system "
        "provides end-to-end automation from contract generation to drift detection, auto-remediation, "
        "and developer notification."
    )
    p = doc.add_paragraph(conclusion_text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    conclusion_text2 = (
        "The key contribution of this project is the integration of multiple automation techniques "
        "into a cohesive pipeline. Rather than treating contract testing as an isolated activity, "
        "the system embeds it into the development workflow such that contract compliance is "
        "continuously verified, issues are automatically detected and fixed, and developers are "
        "immediately notified — all without requiring manual intervention."
    )
    p = doc.add_paragraph(conclusion_text2)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    conclusion_text3 = (
        "The project has been deployed to the team's GitLab CI/CD pipeline and is actively "
        "being used to validate API contracts on every code push. Future enhancements could "
        "include support for additional contract formats (Pact, gRPC), integration with "
        "API gateway metrics for production drift detection, and machine learning-based "
        "contract generation that learns from API usage patterns."
    )
    p = doc.add_paragraph(conclusion_text3)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    conclusion_text4 = (
        "This internship provided invaluable experience in full-stack development, DevOps "
        "practices, and automation engineering. The skills acquired — from Java/Spring Boot "
        "and Python development to CI/CD pipeline design and notification systems — form a "
        "strong foundation for a career in software engineering."
    )
    p = doc.add_paragraph(conclusion_text4)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()

    # ================================================================
    # CHAPTER 5: REFERENCES
    # ================================================================
    add_heading_formatted(doc, "5. References")

    references = [
        '[1] Spring Cloud Contract Documentation, "Spring Cloud Contract Reference Guide", '
        'VMware, 2024. Available: https://docs.spring.io/spring-cloud-contract/',
        '[2] OpenAPI Initiative, "OpenAPI Specification v3.0", The Linux Foundation, 2021. '
        'Available: https://spec.openapis.org/oas/v3.0.3',
        '[3] S. Newman, "Building Microservices: Designing Fine-Grained Systems", '
        "O'Reilly Media, 2nd Edition, 2021.",
        '[4] GitLab Documentation, "GitLab CI/CD Pipeline Configuration Reference", '
        'GitLab Inc., 2024. Available: https://docs.gitlab.com/ee/ci/',
        '[5] Spring Boot Documentation, "Spring Boot Reference Guide v3.2", '
        'VMware, 2024. Available: https://docs.spring.io/spring-boot/docs/3.2.x/reference/',
        '[6] M. Fowler, "ContractTest", Martin Fowler Blog, 2011. '
        'Available: https://martinfowler.com/bliki/ContractTest.html',
        '[7] Python Software Foundation, "Python 3 Documentation", 2024. '
        'Available: https://docs.python.org/3/',
        '[8] Flask Documentation, "Flask Web Framework", Pallets Projects, 2024. '
        'Available: https://flask.palletsprojects.com/',
    ]

    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)

    return doc


def add_heading_formatted(doc, text):
    """Adds a formatted heading (bold, underline, 12pt, left-justified)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    doc.add_paragraph()


def add_subheading_formatted(doc, text):
    """Adds a formatted subheading."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    doc.add_paragraph()


def add_sub_subheading(doc, text):
    """Adds a sub-subheading (bold, no underline)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"


if __name__ == "__main__":
    doc = create_document()
    output_path = r"c:\Projects\contract-testing-automation\Internship_Report_Druva_S_Kumar.docx"
    doc.save(output_path)
    print(f"Report generated: {output_path}")
